import os
from threading import Thread
from flask import Flask

import logging
from uuid import uuid4
from telegram import Update
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler,
    ContextTypes, ConversationHandler, filters
)
from docx import Document
from telegram import ReplyKeyboardMarkup, KeyboardButton

# === Налаштування ===
BOT_TOKEN = os.getenv("BOT_TOKEN")  # встав свій токен
TEMPLATE_PATH = "template.docx"
OUTPUT_DIR = "generated"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Поля для підстановки
FIELDS = [
    "contract_number",      # № договору
    "contract_date",        # дата договору
    "customer",             # замовник
    "amount",               # сума
    "items_total_text"      # всього найменувань текстом
]

# Людські назви полів і підказки
FIELD_PROMPTS = {
    "contract_number": "Введіть номер договору (id.день.місяць.рік.номер_договору):",
    "contract_date": "Введіть дату укладання договору (Приклад: 1 січня 2025р.):",
    "customer": "Введіть замовника (Приклад: ФОП Прізвище Ім'я Побатькові):",
    "amount": "Введіть суму (Приклад: 4000,00):",
    "items_total_text": "Всього найменувань (Приклад: чотири тисячі):"
}

# Валідація типів даних (text / number)
FIELD_TYPES = {
    "contract_number": "mixed",
    "contract_date": "text",
    "customer": "text",
    "amount": "number",
    "items_total_text": "text"
}




# === Логування ===
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# === Стани діалогу ===
ASKING = 1

# === Функції для Word ===
def replace_placeholders_in_paragraph(paragraph, mapping):
    text = paragraph.text
    for key, val in mapping.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in text:
            text = text.replace(placeholder, val)
    if text != paragraph.text:
        for run in paragraph.runs:
            p = run._element
            p.getparent().remove(p)
        paragraph.add_run(text)

def replace_placeholders_in_docx(doc_path, mapping, output_path):
    doc = Document(doc_path)
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, mapping)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, mapping)
    doc.save(output_path)

# === Обробники Telegram ===
async def start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data.clear()

    keyboard = [
        [KeyboardButton("Сформувати рахунок📋")],
        [KeyboardButton("Скасувати🔸")]
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    await update.message.reply_text(
        "Привіт! 👋\n"
        "Я створюю рахунок на оплату за шаблоном Word.\n"
        "Натисни кнопку нижче, щоб почати або вийти.",
        reply_markup=reply_markup
    )

async def cancel(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data.clear()  # очищає дані користувача
    keyboard = [
        [KeyboardButton("Сформувати рахунок📋")],
        [KeyboardButton("Скасувати🔸")]
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    await update.message.reply_text(
        "✅ Скасовано. Ви можете почати знову, натиснувши 'Сформувати рахунок'.",
        reply_markup=reply_markup
    )
    return ConversationHandler.END


async def form_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data["form_values"] = {}
    ctx.user_data["form_index"] = 0
    first_field = FIELDS[0]
    await update.message.reply_text(FIELD_PROMPTS[first_field])
    return ASKING

async def form_ask(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()

    # Якщо користувач натиснув кнопку "Скасувати🔸" як текст
    if text == "Скасувати🔸" or text.lower() in ["скасувати", "вийти", "cancel", "stop"]:
        return await cancel(update, ctx)

    idx = ctx.user_data.get("form_index", 0)
    field = FIELDS[idx]

    # Перевірка типу даних
    field_type = FIELD_TYPES.get(field, "text")

    if field_type == "number":
        import re
        if not re.match(r'^[0-9]+([.,][0-9]+)?$', text):
            await update.message.reply_text("⚠️ Введіть, будь ласка, лише число (наприклад: 4000,00):")
            return ASKING

    elif field_type == "text":
        # дозволяємо будь-який текст, але відкидаємо чисто цифрові рядки
        if text.replace(',', '').replace('.', '').isdigit():
            await update.message.reply_text("⚠️ Це має бути текст, а не число. Спробуйте ще раз:")
            return ASKING

    elif field_type == "mixed":
        import re
        # дозволяємо цифри, букви, крапки, дефіси і підкреслення
        if not re.match(r'^[0-9A-Za-zА-Яа-яіїєІЇЄ.\-_/]+$', text):
            await update.message.reply_text("⚠️ Невірний формат. Використовуйте лише цифри, літери та крапки/дефіси.")
            return ASKING

    # Записуємо значення (тільки один раз)
    ctx.user_data["form_values"][field] = text
    idx += 1

    if idx < len(FIELDS):
        ctx.user_data["form_index"] = idx
        next_field = FIELDS[idx]
        # --> ось тут використовуємо дружній текст з FIELD_PROMPTS
        prompt = FIELD_PROMPTS.get(next_field, f"Введіть {next_field}:")
        await update.message.reply_text(prompt)
        return ASKING
    else:
        mapping = ctx.user_data["form_values"]
        out_name = f"{uuid4().hex}.docx"
        out_path = os.path.join(OUTPUT_DIR, out_name)
        try:
            replace_placeholders_in_docx(TEMPLATE_PATH, mapping, out_path)
        except Exception as e:
            logger.exception("Помилка при формуванні документа:")
            await update.message.reply_text("⚠️ Виникла помилка при створенні файлу.")
            return ConversationHandler.END

        contract_number = mapping.get("contract_number", "без_номеру").replace(":", ".").replace("/", ".").replace("\\",
                                                                                                                   ".")

        # Формуємо гарну назву файлу
        file_name = f"№{contract_number} Рахунок на оплату.docx"

        await update.message.reply_document(open(out_path, "rb"), filename=file_name)
        await update.message.reply_text(f"✅ Готово! Документ створено:\n📄 {file_name}")

        return ConversationHandler.END


async def cancel(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Скасовано ✅")
    return ConversationHandler.END

def main():
    app = ApplicationBuilder().token(BOT_TOKEN).build()

    conv = ConversationHandler(
        entry_points=[
            CommandHandler("form", form_start),
            MessageHandler(filters.Regex("^Сформувати рахунок📋$"), form_start)
        ],
        states={
            ASKING: [MessageHandler(filters.TEXT & ~filters.COMMAND, form_ask)]
        },
        fallbacks=[
            CommandHandler("cancel", cancel),
            MessageHandler(filters.Regex("^Скасувати🔸$"), cancel)
        ]
    )

    app.add_handler(CommandHandler("start", start))
    app.add_handler(conv)

    print("🤖 Бот запущено...")
    app.run_polling()

if __name__ == "__main__":
    main()
